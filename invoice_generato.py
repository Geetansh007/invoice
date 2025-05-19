from docx import Document
import os

class InvoiceDocGenerator:
    def __init__(self, template_path: str, output_dir: str, data_list: list):
        self.template_path = template_path
        self.output_dir = output_dir
        self.data_list = data_list

        os.makedirs(self.output_dir, exist_ok=True)

        # Placeholder-to-field mapping
        self.mapping = {
            "Col A": "Date of Entry",
            "Col B": "Amount",
            "Col C": "Description",
            "Col D": "Beneficiary Name",
            "Col E": "Bank Name",
            "Col F": "Bank Account No",
            "Col G": "IFSC/SWIFT Code",
            "Col H": "IBAN No",
            "Col I": "Address"
        }

    def generate_documents(self):
        for data in self.data_list:
            doc = Document(self.template_path)
            self._replace_placeholders(doc, data)
            self._replace_self_generates(doc, data)

            filename = f"{data.get('Beneficiary Name', 'Unknown')}.docx"
            output_path = os.path.join(self.output_dir, filename)
            doc.save(output_path)
            print(f"Saved: {output_path}")

    def _replace_placeholders(self, doc, data):
        for para in doc.paragraphs:
            self._replace_text_in_para(para, data)

        for table in doc.tables:
            if self._is_invoice_item_table(table):
                self._fill_invoice_rows_in_place(table, data)
            else:
                for row in table.rows:
                    for cell in row.cells:
                        self._replace_text_in_para(cell, data)

    def _replace_text_in_para(self, container, data):
        for placeholder, key in self.mapping.items():
            value = str(data.get(key, ""))
            if placeholder in container.text:
                container.text = container.text.replace(placeholder, value)

    def _is_invoice_item_table(self, table):
        # Identify table by checking headers
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        expected = ['sr. no.', 'description', 'unit', 'unit/ price', 'amount']
        return all(any(header == e for header in headers) for e in expected)

    def _fill_invoice_rows_in_place(self, table, data):
        descriptions = data.get("Description", [])
        amounts = data.get("Amount", [])

        if not isinstance(descriptions, list):
            descriptions = [descriptions]
        if not isinstance(amounts, list):
            amounts = [amounts]

        if len(descriptions) != len(amounts):
            print("Mismatch between Description and Amount lists.")
            return

        row_index = 1  # skip header row
        for i, (desc, amt) in enumerate(zip(descriptions, amounts), start=1):
            if row_index >= len(table.rows):
                print("Not enough rows available in the template to insert all items.")
                break

            row = table.rows[row_index]
            cells = row.cells
            if len(cells) < 5:
                continue  

            cells[0].text = str(i)           
            cells[1].text = str(desc)                   
            cells[-1].text = str(amt)       
            row_index += 1

    def _replace_self_generates(self, doc, data):
        count = 0

        def replace_in_text(text):
            nonlocal count
            if "Self generate" not in text:
                return text

            parts = text.split("Self generate")
            result = parts[0]
            for i in range(1, len(parts)):
                count += 1
                if count == 1:
                    replacement = data.get("Invoice Number", "")
                elif count == 3:
                    replacement = str(data.get("Total Amount", ""))
                elif count == 4:
                    replacement = data.get("Total Amount in Words", "")
                else:
                    replacement = "Self generate"
                result += replacement + parts[i]
            return result

        for para in doc.paragraphs:
            para.text = replace_in_text(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = replace_in_text(cell.text)
